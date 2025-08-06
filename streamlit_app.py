# --- Install Required Packages ---
# Install these using:
# pip install streamlit requests python-docx transformers

import streamlit as st
import requests
from io import BytesIO
from docx import Document
import random
import re

# Hugging Face inference API
import json
import time

# --- CONFIG ---
st.set_page_config(page_title="AI Research Assistant (Open Source)", layout="wide")
st.title("📚 Open-Source AI Research Assistant with Real Citations")

# --- Sidebar ---
st.sidebar.header("🔧 Settings")
topic = st.sidebar.text_input("Research Topic or Question")
citation_count = st.sidebar.slider("Number of citations", 2, 10, 4)
citation_style = st.sidebar.selectbox("Citation Style", ["APA", "MLA"])
sections_to_generate = st.sidebar.multiselect(
    "Sections to generate",
    ["Abstract", "Introduction", "Problem Statement", "Methodology", "Conclusion"],
    default=["Introduction"]
)

# --- Helper Functions ---
def fetch_citations(topic, n):
    url = f"https://api.semanticscholar.org/graph/v1/paper/search?query={topic}&limit={n}&fields=title,authors,year,url"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json().get("data", [])
        citations = []
        for paper in data:
            authors_list = [a["name"] for a in paper["authors"]][:2]
            authors = ", ".join(authors_list)
            year = paper.get("year", "n.d.")
            title = paper.get("title", "No Title")
            url = paper.get("url", "")
            if citation_style == "APA":
                citation = f"({authors}, {year})"
                ref = f"{authors} ({year}). {title}. Available at: {url}"
            else:
                citation = f"({authors})"
                ref = f"{authors}. \"{title}.\" {year}, {url}."
            citations.append({
                "citation": citation,
                "ref": ref
            })
        return citations
    else:
        return []

# Hugging Face Inference API for text generation
def hf_generate(topic, section):
    prompt = f"""
Write a 5-sentence {section.lower()} for a research paper titled: \"{topic}\".
Use formal academic language. Do not include citations.
"""
    headers = {"Authorization": "Bearer hf_dummytoken"}  # Replace with actual token if needed
    json_data = {
        "inputs": prompt,
        "parameters": {"temperature": 0.7, "max_new_tokens": 300}
    }
    response = requests.post(
        "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.1",
        headers=headers, json=json_data
    )
    if response.status_code == 200:
        try:
            return response.json()[0]["generated_text"].strip().split("\n")[-1]
        except:
            return "[Error parsing model output]"
    else:
        return f"[Error from HuggingFace API: {response.status_code}]"

def insert_citations(text, citations):
    sentences = re.split(r'(?<=[.!?]) +', text)
    selected = random.sample(citations, min(len(citations), len(sentences)))
    result = ""
    for i, sentence in enumerate(sentences):
        sentence = sentence.strip()
        if sentence:
            citation = selected[i % len(selected)]["citation"]
            result += f"{sentence} {citation} "
    return result.strip()

# --- Main Generation ---
if st.button("🚀 Generate Selected Sections"):
    if not topic:
        st.warning("Please enter your research topic.")
    else:
        final_sections = {}
        all_references = []

        with st.spinner("Generating content and fetching citations..."):
            for section in sections_to_generate:
                raw_text = hf_generate(topic, section)
                citations = fetch_citations(topic, citation_count)
                if citations:
                    content = insert_citations(raw_text, citations)
                    final_sections[section] = content
                    all_references.extend([c["ref"] for c in citations])
                else:
                    st.error(f"Failed to fetch citations for {section}.")

        st.success("✅ All sections generated!")

        for section, text in final_sections.items():
            st.subheader(section)
            st.write(text)

        st.subheader("📚 References")
        for ref in all_references:
            st.markdown(f"- {ref}")

        # --- Export ---
        st.markdown("---")
        st.subheader("📥 Download Full Document")

        title = st.text_input("Enter Document Title", value=topic.title())

        if st.button("Download DOCX"):
            doc = Document()
            doc.add_heading(title, 0)
            for section, text in final_sections.items():
                doc.add_heading(section, level=1)
                doc.add_paragraph(text)
            doc.add_heading("References", level=1)
            for ref in all_references:
                doc.add_paragraph(ref)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("Download .docx", data=buffer, file_name="research_paper.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
